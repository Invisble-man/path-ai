import streamlit as st
import pandas as pd
import sqlite3
import json
import re
import io
import zipfile
import base64
import urllib.request
from dataclasses import dataclass, asdict
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any

from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill

import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================================================
# CONFIG
# =========================================================
APP_NAME = "Path.ai"
APP_VERSION = "v1.5.0"
DB_PATH = "path_state.db"

PAGES = ["Task List", "Intake", "Company", "Compliance", "Draft", "Export"]

GL_ACTIONABLE = "ACTIONABLE"
GL_INFORMATIONAL = "INFORMATIONAL"
GL_IRRELEVANT = "IRRELEVANT"
GL_AUTO = "AUTO_RESOLVED"

S_PASS = "pass"
S_FAIL = "fail"
S_UNKNOWN = "unknown"

BUCKETS = [
    "Top Priorities",
    "Blockers",
    "Submission & Format",
    "Required Forms & Registrations",
    "Attachments/Exhibits",
    "Technical Requirements",
    "Pricing & Cost",
    "Past Performance",
    "Other",
]

ACTIONABLE_STRONG = [
    "must", "shall", "required", "offeror shall", "submit", "provide", "complete",
    "include", "fill", "deliver", "due", "deadline", "no later than",
    "format", "font", "margin", "page limit", "pages", "electronic", "pdf", "excel",
    "spreadsheet", "sf1449", "sf-1449", "block", "volume", "technical", "price", "cost",
    "pricing", "attachment", "exhibit", "forms", "representations", "certifications",
    "sam.gov", "uei", "cage"
]

INFORMATIONAL_HINTS = [
    "background", "purpose", "overview", "the government", "will", "may", "should",
    "intended", "general", "note:", "reference", "definitions", "acronyms"
]

IRRELEVANT_POST_AWARD = [
    "invoice", "invoicing", "payment", "paid", "warranty", "claims", "disputes",
    "contractor shall bill", "final invoice", "prompt payment", "modification",
    "change request", "deobligate", "termination for convenience"
]

AUTO_RESOLVE_HINTS = [
    "not applicable", "n/a", "none required", "no action required"
]

CRITICAL_PATTERNS = {
    "Submission deadline": [
        r"offer due date",
        r"proposal(?:s)?\s+due",
        r"due\s+date",
        r"deadline",
        r"no later than"
    ],
    "Submission method": [
        r"submit\s+electronically",
        r"email\s+to",
        r"via\s+.*portal",
        r"upload",
        r"submit(?:tal)?\s+through"
    ],
    "File format rules": [
        r"\bpdf\b",
        r"editable\s+spreadsheet",
        r"\bexcel\b",
        r"file\s+format",
        r"\bfont\b",
        r"\bmargin\b",
        r"page\s+limit"
    ],
    "Required forms (SF/attachments)": [
        r"sf\s*1449",
        r"sf-\s*1449",
        r"\bsf\s*\d{3,5}\b",
        r"block\s+\d+",
        r"representations?\s+and\s+certifications?",
        r"reps?\s*&\s*certs?"
    ],
    "Required attachments/exhibits": [
        r"attachment\s+[a-z0-9]+",
        r"exhibit\s+[a-z0-9]+",
        r"include\s+the\s+following",
        r"submit\s+the\s+following"
    ],
}


# =========================================================
# DATA MODEL
# =========================================================
@dataclass
class Item:
    id: str
    raw_text: str
    normalized_text: str
    source: str
    section: str
    bucket: str
    gating_label: str
    confidence: float
    status: str
    done: bool
    mapped_section: str
    notes: str

    # traceability
    source_file: str
    page_number: Optional[int]
    source_snippet: str

    is_critical: bool
    created_at: str


# =========================================================
# DB LAYER
# =========================================================
def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS runs (
            run_id TEXT PRIMARY KEY,
            name TEXT,
            created_at TEXT,
            updated_at TEXT,
            intake_json TEXT,
            company_json TEXT,
            workflow_json TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS items (
            run_id TEXT,
            item_id TEXT,
            json TEXT,
            PRIMARY KEY (run_id, item_id)
        )
    """)
    conn.commit()
    return conn


def now_iso() -> str:
    return datetime.utcnow().isoformat()


def create_run(run_name: str) -> str:
    run_id = f"run_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"
    now = now_iso()
    conn = db()
    conn.execute(
        "INSERT INTO runs (run_id, name, created_at, updated_at, intake_json, company_json, workflow_json) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            run_id, run_name, now, now,
            json.dumps({
                "sol_text": "",
                "sources": [],
                "diagnostics": {},
                "drafts": {}
            }),
            json.dumps({
                "legal_name": "",
                "uei": "",
                "cage": "",
                "naics": "",
                "address": "",
                "primary_contact": "",
                "email": "",
                "phone": "",
                "sam_registered": False,
                "certs": [],
                "contract_title": "",
                "logo_b64": ""
            }),
            json.dumps({
                "intake_complete": False,
                "company_complete": False,
                "compliance_complete": False,
                "draft_complete": False
            })
        )
    )
    conn.commit()
    return run_id


def list_runs() -> List[Tuple[str, str, str]]:
    conn = db()
    return conn.execute("SELECT run_id, name, updated_at FROM runs ORDER BY updated_at DESC").fetchall()


def load_run_meta(run_id: str) -> Dict[str, Any]:
    conn = db()
    row = conn.execute(
        "SELECT intake_json, company_json, workflow_json, name, created_at, updated_at FROM runs WHERE run_id=?",
        (run_id,)
    ).fetchone()
    if not row:
        return {"intake": {}, "company": {}, "workflow": {}, "name": "", "created_at": "", "updated_at": ""}
    intake_json, company_json, workflow_json, name, created_at, updated_at = row
    return {
        "intake": json.loads(intake_json or "{}"),
        "company": json.loads(company_json or "{}"),
        "workflow": json.loads(workflow_json or "{}"),
        "name": name,
        "created_at": created_at,
        "updated_at": updated_at
    }


def save_run_meta(run_id: str, intake: Dict, company: Dict, workflow: Dict, run_name: Optional[str] = None):
    current = load_run_meta(run_id)
    name_to_save = run_name if run_name is not None else current.get("name", "")
    conn = db()
    conn.execute(
        "UPDATE runs SET intake_json=?, company_json=?, workflow_json=?, name=?, updated_at=? WHERE run_id=?",
        (json.dumps(intake), json.dumps(company), json.dumps(workflow), name_to_save, now_iso(), run_id)
    )
    conn.commit()


def _item_defaults_patch(d: Dict[str, Any]) -> Dict[str, Any]:
    patched = dict(d)
    patched.setdefault("id", patched.get("item_id", patched.get("id", "")) or patched.get("id", ""))
    patched.setdefault("raw_text", patched.get("text", patched.get("normalized_text", "")) or "")
    patched.setdefault("normalized_text", patched.get("text", patched.get("raw_text", "")) or "")
    patched.setdefault("source", "RFP")
    patched.setdefault("section", "General")
    patched.setdefault("bucket", "Other")
    patched.setdefault("gating_label", GL_INFORMATIONAL)
    patched.setdefault("confidence", 0.50)
    patched.setdefault("status", S_UNKNOWN)
    patched.setdefault("done", False)
    patched.setdefault("mapped_section", "General / Supporting")
    patched.setdefault("notes", "")

    patched.setdefault("source_file", patched.get("source", "Unknown"))
    patched.setdefault("page_number", None)
    patched.setdefault("source_snippet", (patched.get("normalized_text", "") or "")[:220])

    patched.setdefault("is_critical", False)
    patched.setdefault("created_at", now_iso())

    for legacy in ["item_id", "text"]:
        patched.pop(legacy, None)
    return patched


def load_items(run_id: str) -> Dict[str, Item]:
    conn = db()
    rows = conn.execute("SELECT item_id, json FROM items WHERE run_id=?", (run_id,)).fetchall()
    out: Dict[str, Item] = {}
    for item_id, j in rows:
        d = _item_defaults_patch(json.loads(j))
        try:
            out[item_id] = Item(**d)
        except Exception:
            out[item_id] = Item(
                id=item_id,
                raw_text=str(d.get("raw_text", "")),
                normalized_text=str(d.get("normalized_text", "")),
                source=str(d.get("source", "RFP")),
                section=str(d.get("section", "General")),
                bucket=str(d.get("bucket", "Other")),
                gating_label=str(d.get("gating_label", GL_INFORMATIONAL)),
                confidence=float(d.get("confidence", 0.5)),
                status=str(d.get("status", S_UNKNOWN)),
                done=bool(d.get("done", False)),
                mapped_section=str(d.get("mapped_section", "General / Supporting")),
                notes=str(d.get("notes", "")),
                source_file=str(d.get("source_file", "Unknown")),
                page_number=d.get("page_number", None),
                source_snippet=str(d.get("source_snippet", ""))[:220],
                is_critical=bool(d.get("is_critical", False)),
                created_at=str(d.get("created_at", now_iso()))
            )
    return out


def upsert_item(run_id: str, item: Item):
    conn = db()
    conn.execute(
        "INSERT OR REPLACE INTO items (run_id, item_id, json) VALUES (?, ?, ?)",
        (run_id, item.id, json.dumps(asdict(item)))
    )
    conn.commit()


def delete_all_items(run_id: str):
    conn = db()
    conn.execute("DELETE FROM items WHERE run_id=?", (run_id,))
    conn.commit()


# =========================================================
# TEXT / CLASSIFICATION
# =========================================================
def clean_text(t: str) -> str:
    t = (t or "").strip()
    t = re.sub(r"\s+", " ", t)
    return t


def normalize_text(t: str) -> str:
    t = clean_text(t)
    t = re.sub(r"(?i)\battachment mention:\s*", "", t)
    t = re.sub(r"(?i)\bmapped section:\s*", "", t)
    return t


def infer_section(line: str) -> str:
    l = line.lower()
    if any(k in l for k in ["deadline", "due date", "no later than", "submit", "submission", "format", "font", "margin", "page limit"]):
        return "Submission"
    if any(k in l for k in ["sf-", "sf1449", "representations", "certifications", "sam.gov", "uei", "cage"]):
        return "Forms & Registration"
    if any(k in l for k in ["past performance", "experience", "references"]):
        return "Past Performance"
    if any(k in l for k in ["price", "pricing", "cost", "rates", "clins", "labor categories"]):
        return "Pricing"
    if any(k in l for k in ["technical", "approach", "method", "pws", "sow", "performance work statement"]):
        return "Technical"
    if any(k in l for k in ["attachment", "exhibit", "appendix"]):
        return "Attachments"
    return "General"


def mapped_section_from_section(section: str) -> str:
    mapping = {
        "Submission": "Volume 1 – Submission / Admin",
        "Forms & Registration": "Forms / Reps & Certs",
        "Technical": "Volume 2 – Technical Approach",
        "Pricing": "Volume 3 – Pricing",
        "Past Performance": "Volume 4 – Past Performance",
        "Attachments": "Attachments / Exhibits",
        "General": "General / Supporting",
    }
    return mapping.get(section, "General / Supporting")


def bucketize(text: str, section: str) -> str:
    t = (text or "").lower()
    s = (section or "").lower()
    if "submission" in s or any(k in t for k in ["deadline", "due", "format", "font", "margin", "page limit", "pdf", "excel", "electronic"]):
        return "Submission & Format"
    if any(k in t for k in ["sf1449", "sf-1449", "representations", "certifications", "sam.gov", "uei", "cage"]):
        return "Required Forms & Registrations"
    if any(k in t for k in ["past performance", "references", "experience"]):
        return "Past Performance"
    if any(k in t for k in ["price", "pricing", "cost", "rates", "labor category", "clins"]):
        return "Pricing & Cost"
    if any(k in t for k in ["technical", "approach", "method", "pws", "sow"]):
        return "Technical Requirements"
    if any(k in t for k in ["attachment", "exhibit", "appendix"]):
        return "Attachments/Exhibits"
    return "Other"


def classify_gating(text: str) -> str:
    t = (text or "").lower()

    if any(h in t for h in AUTO_RESOLVE_HINTS):
        return GL_AUTO

    if any(h in t for h in IRRELEVANT_POST_AWARD):
        if "proposal" in t or "offer" in t:
            return GL_ACTIONABLE
        return GL_IRRELEVANT

    hits = sum(1 for k in ACTIONABLE_STRONG if k in t)
    if hits >= 2:
        return GL_ACTIONABLE

    if any(h in t for h in INFORMATIONAL_HINTS):
        return GL_INFORMATIONAL

    return GL_INFORMATIONAL


def confidence_score(text: str, gating: str) -> float:
    base = 0.55
    t = (text or "").lower()
    actionable_hits = sum(1 for k in ACTIONABLE_STRONG if k in t)
    info_hits = sum(1 for k in INFORMATIONAL_HINTS if k in t)

    if gating == GL_ACTIONABLE:
        base = min(0.95, 0.55 + actionable_hits * 0.07)
    elif gating == GL_INFORMATIONAL:
        base = min(0.85, 0.45 + info_hits * 0.08)
    elif gating == GL_AUTO:
        base = 0.85
    elif gating == GL_IRRELEVANT:
        base = 0.80

    return float(max(0.30, min(0.99, base)))


def detect_critical(text: str) -> bool:
    t = (text or "").lower()
    critical_keywords = [
        "deadline", "due", "no later than", "submit", "submission", "file format",
        "pdf", "excel", "font", "margin", "page limit", "sf1449", "sf-1449", "block",
        "attachment", "exhibit", "representations", "certifications"
    ]
    return any(k in t for k in critical_keywords)


def dedupe_items(items: List[Item]) -> List[Item]:
    seen: Dict[str, Item] = {}
    priority = {GL_ACTIONABLE: 3, GL_INFORMATIONAL: 2, GL_AUTO: 1, GL_IRRELEVANT: 0}
    for it in items:
        key = re.sub(r"[^a-z0-9 ]", "", it.normalized_text.lower())
        key = re.sub(r"\s+", " ", key).strip()
        if key in seen:
            if priority[it.gating_label] > priority[seen[key].gating_label]:
                seen[key] = it
            continue
        seen[key] = it
    return list(seen.values())


def build_missing_critical_tasks(sol_text: str) -> List[Item]:
    t = (sol_text or "").lower()
    tasks: List[Item] = []
    created = now_iso()

    for name, patterns in CRITICAL_PATTERNS.items():
        found = False
        for p in patterns:
            if re.search(p, t, re.IGNORECASE):
                found = True
                break
        if not found:
            item_id = f"crit_{re.sub(r'[^a-z0-9]+', '_', name.lower()).strip('_')}"
            section = "Critical Fields"
            bucket = "Submission & Format" if "submission" in name.lower() or "format" in name.lower() else "Required Forms & Registrations"
            tasks.append(Item(
                id=item_id,
                raw_text=f"Confirm/enter: {name}",
                normalized_text=f"Confirm/enter: {name}",
                source="Detected",
                section=section,
                bucket=bucket,
                gating_label=GL_ACTIONABLE,
                confidence=0.90,
                status=S_UNKNOWN,
                done=False,
                mapped_section=mapped_section_from_section("Submission"),
                notes="This field was not detected. Verify in the RFP and enter it manually.",
                source_file="Detected",
                page_number=None,
                source_snippet="Not detected in extracted text. Verify directly in the solicitation.",
                is_critical=True,
                created_at=created
            ))
    return tasks


# =========================================================
# DIAGNOSTICS + PDF EXTRACTION (like your screenshot)
# =========================================================
def extraction_quality_label(total_pages: int, pages_with_text: int, chars: int) -> str:
    if total_pages <= 0:
        return "Unknown"
    if pages_with_text == 0 or chars < 300:
        return "Poor"
    ratio = pages_with_text / max(1, total_pages)
    chars_per_page = chars / max(1, pages_with_text)

    if ratio >= 0.90 and chars_per_page >= 1500:
        return "Excellent"
    if ratio >= 0.75 and chars_per_page >= 800:
        return "Good"
    if ratio >= 0.50 and chars_per_page >= 400:
        return "Fair"
    return "Poor"


def likely_scanned(total_pages: int, pages_with_text: int, chars: int) -> bool:
    if total_pages <= 0:
        return False
    ratio = pages_with_text / max(1, total_pages)
    if ratio < 0.25:
        return True
    if chars < max(800, total_pages * 120):
        return True
    return False


def extract_pages_from_pdfs(files: List) -> Tuple[List[Dict[str, Any]], List[str], Dict[str, Any]]:
    pages_data: List[Dict[str, Any]] = []
    sources: List[str] = []
    total_pages = 0
    pages_with_text = 0
    chars = 0

    for f in files:
        fname = getattr(f, "name", "uploaded.pdf")
        try:
            reader = PdfReader(f)
            pages = len(reader.pages)
            total_pages += pages
            sources.append(f"{fname} ({pages} pages)")
            for i in range(pages):
                txt = reader.pages[i].extract_text() or ""
                txt = txt.strip()
                if txt:
                    pages_with_text += 1
                    chars += len(txt)
                    pages_data.append({
                        "file_name": fname,
                        "page_number": i + 1,
                        "text": txt
                    })
        except Exception:
            sources.append(f"{fname} (0 pages — extraction failed)")

    diag = {
        "file_type": "pdf",
        "pages": total_pages,
        "pages_with_text": pages_with_text,
        "characters_extracted": chars,
        "likely_scanned": "Yes" if likely_scanned(total_pages, pages_with_text, chars) else "No",
        "extraction_quality": extraction_quality_label(total_pages, pages_with_text, chars)
    }
    return pages_data, sources, diag


def split_candidates_from_page(text: str) -> List[str]:
    raw_lines = [clean_text(x) for x in (text or "").splitlines()]
    raw_lines = [x for x in raw_lines if x and len(x) >= 10]

    candidates: List[str] = []
    for ln in raw_lines:
        lnl = ln.lower()
        if re.match(r"^\d+\s*/\s*\d+$", ln):
            continue
        if lnl.startswith("page ") and len(ln) < 20:
            continue

        parts = re.split(r"(?<=[.;:])\s+(?=[A-Z0-9])", ln)
        for p in parts:
            p = clean_text(p)
            if p and len(p) >= 12:
                candidates.append(p)

    seen = set()
    out = []
    for c in candidates:
        key = c.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
    return out


def build_items_from_sources(
    combined_text: str,
    pages_data: List[Dict[str, Any]],
    pasted_text: str,
    include_reference: bool
) -> List[Item]:
    built: List[Item] = []
    created = now_iso()

    # missing critical tasks first
    built.extend(build_missing_critical_tasks(combined_text))

    idx = 1

    # PDF-derived
    for page in pages_data:
        fname = page["file_name"]
        pno = page["page_number"]
        page_text = page["text"]
        candidates = split_candidates_from_page(page_text)

        for ln in candidates:
            raw = clean_text(ln)
            if len(raw) < 12:
                continue

            gating = classify_gating(raw)
            if gating == GL_IRRELEVANT:
                continue
            if gating == GL_INFORMATIONAL and not include_reference:
                continue

            section = infer_section(raw)
            mapped = mapped_section_from_section(section)
            bucket = bucketize(raw, section)
            conf = confidence_score(raw, gating)

            done = False
            status = S_UNKNOWN
            if gating == GL_AUTO:
                done = True
                status = S_PASS

            item_id = f"r{idx:05d}_{re.sub(r'[^a-z0-9]+','_', raw.lower())[:16].strip('_')}"
            idx += 1

            built.append(Item(
                id=item_id,
                raw_text=raw,
                normalized_text=normalize_text(raw),
                source="RFP",
                section=section,
                bucket=bucket,
                gating_label=gating,
                confidence=float(conf),
                status=status,
                done=done,
                mapped_section=mapped,
                notes="",
                source_file=fname,
                page_number=pno,
                source_snippet=normalize_text(raw)[:220],
                is_critical=detect_critical(raw) or item_id.startswith("crit_"),
                created_at=created
            ))

    # Pasted text-derived
    if pasted_text and pasted_text.strip():
        lines = [clean_text(x) for x in pasted_text.splitlines() if clean_text(x)]
        for ln in lines:
            raw = clean_text(ln)
            if len(raw) < 12:
                continue

            gating = classify_gating(raw)
            if gating == GL_IRRELEVANT:
                continue
            if gating == GL_INFORMATIONAL and not include_reference:
                continue

            section = infer_section(raw)
            mapped = mapped_section_from_section(section)
            bucket = bucketize(raw, section)
            conf = confidence_score(raw, gating)

            done = False
            status = S_UNKNOWN
            if gating == GL_AUTO:
                done = True
                status = S_PASS

            item_id = f"t{idx:05d}_{re.sub(r'[^a-z0-9]+','_', raw.lower())[:16].strip('_')}"
            idx += 1

            built.append(Item(
                id=item_id,
                raw_text=raw,
                normalized_text=normalize_text(raw),
                source="RFP",
                section=section,
                bucket=bucket,
                gating_label=gating,
                confidence=float(conf),
                status=status,
                done=done,
                mapped_section=mapped,
                notes="",
                source_file="Pasted text",
                page_number=None,
                source_snippet=normalize_text(raw)[:220],
                is_critical=detect_critical(raw) or item_id.startswith("crit_"),
                created_at=created
            ))

    return built


# =========================================================
# COMPANY-BASED AUTO-RESOLVE
# =========================================================
def company_auto_resolve(items: Dict[str, Item], company: Dict[str, Any]) -> Tuple[Dict[str, Item], int]:
    uei = clean_text(company.get("uei", ""))
    cage = clean_text(company.get("cage", ""))
    sam = bool(company.get("sam_registered", False))
    certs = company.get("certs", []) or []
    certs_lower = set([str(c).lower() for c in certs])

    changed = 0
    updated = dict(items)

    def should_autoresolve(text: str) -> bool:
        t = (text or "").lower()

        if any(k in t for k in ["sam.gov", "sam registration", "registered in sam"]):
            return sam

        if "uei" in t and uei:
            return True
        if "cage" in t and cage:
            return True

        if any(k in t for k in ["sdvosb", "service-disabled veteran-owned"]):
            return ("sdvosb" in certs_lower) or ("vetcert" in certs_lower) or ("sba vetcert" in certs_lower)

        if "hubzone" in t:
            return "hubzone" in certs_lower
        if "8(a)" in t or "8a" in t:
            return "8(a)" in certs_lower or "8a" in certs_lower
        if "wosb" in t:
            return "wosb" in certs_lower

        return False

    for k, it in items.items():
        if it.gating_label not in [GL_ACTIONABLE, GL_INFORMATIONAL]:
            continue
        if should_autoresolve(it.normalized_text):
            if it.gating_label != GL_AUTO or it.status != S_PASS or not it.done:
                it.gating_label = GL_AUTO
                it.status = S_PASS
                it.done = True
                it.notes = (it.notes + " " if it.notes else "") + "Auto-resolved from Company profile."
                changed += 1
            updated[k] = it

    return updated, changed


# =========================================================
# KPI / GATE / LOCKED FLOW
# =========================================================
def compute_kpis(items: Dict[str, Item]) -> Dict[str, Any]:
    actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
    if not actionable:
        return {
            "completion": 0.0,
            "pass": 0, "fail": 0, "unknown": 0,
            "gate": "WAITING",
            "missing_critical": 0,
            "actionable_total": 0,
        }

    completed = [i for i in actionable if (i.done or i.status in [S_PASS, S_FAIL])]
    completion = len(completed) / max(1, len(actionable))

    pass_ct = sum(1 for i in actionable if i.status == S_PASS)
    fail_ct = sum(1 for i in actionable if i.status == S_FAIL)
    unk_ct = sum(1 for i in actionable if i.status == S_UNKNOWN)

    missing_critical = sum(1 for i in actionable if i.is_critical and i.status == S_UNKNOWN and not i.done)

    if fail_ct > 0:
        gate = "FAIL"
    elif missing_critical > 0:
        gate = "AT RISK"
    elif unk_ct > max(2, int(0.10 * len(actionable))):
        gate = "AT RISK"
    else:
        gate = "PASS"

    return {
        "completion": float(completion),
        "pass": int(pass_ct),
        "fail": int(fail_ct),
        "unknown": int(unk_ct),
        "gate": gate,
        "missing_critical": int(missing_critical),
        "actionable_total": len(actionable),
    }


def workflow_unlocks(run_meta: Dict[str, Any], items: Dict[str, Item]) -> Dict[str, bool]:
    workflow = run_meta.get("workflow", {}) or {}
    intake_ok = bool(workflow.get("intake_complete", False))
    company_ok = bool(workflow.get("company_complete", False))
    compliance_ok = bool(workflow.get("compliance_complete", False))
    draft_ok = bool(workflow.get("draft_complete", False))

    kpis = compute_kpis(items)
    if intake_ok and company_ok and kpis["actionable_total"] > 0:
        if (kpis["completion"] >= 0.60) and (kpis["gate"] != "FAIL"):
            compliance_ok = True
            workflow["compliance_complete"] = True

    run_meta["workflow"] = workflow

    return {
        "Task List": True,
        "Intake": True,
        "Company": intake_ok,
        "Compliance": intake_ok and company_ok,
        "Draft": intake_ok and company_ok and compliance_ok,
        "Export": intake_ok and company_ok and compliance_ok and draft_ok,
    }


# =========================================================
# PRIORITY ENGINE (Top Priorities + Blockers)
# =========================================================
def priority_score(it: Item) -> int:
    t = (it.normalized_text or "").lower()
    score = 0

    if any(k in t for k in ["deadline", "due date", "no later than", "due", "submit by"]):
        score += 60
    if any(k in t for k in ["page limit", "font", "margin", "format", "pdf", "excel", "electronic submission"]):
        score += 45
    if any(k in t for k in ["sf1449", "sf-1449", "sam.gov", "uei", "cage", "representations", "certifications"]):
        score += 40
    if any(k in t for k in ["attachment", "exhibit", "appendix", "include the following"]):
        score += 35
    if any(k in t for k in ["price", "pricing", "cost", "rates", "clins"]):
        score += 28
    if "past performance" in t or "references" in t:
        score += 22
    if any(k in t for k in ["technical", "approach", "method", "pws", "sow"]):
        score += 18

    if it.is_critical:
        score += 35
    if it.status == S_FAIL:
        score += 100
    if it.done or it.status == S_PASS:
        score -= 20

    return score


def blockers(items: Dict[str, Item]) -> List[Item]:
    actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
    out = []
    for it in actionable:
        if it.status == S_FAIL:
            out.append(it)
        elif it.is_critical and (it.status == S_UNKNOWN) and (not it.done):
            out.append(it)
    out = sorted(out, key=lambda x: (-priority_score(x), -x.confidence))
    return out[:12]


def top_priorities(items: Dict[str, Item]) -> List[Item]:
    actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
    actionable = [i for i in actionable if not (i.done or i.status == S_PASS)]
    ranked = sorted(actionable, key=lambda x: (-priority_score(x), -x.confidence))
    return ranked[:10]


# =========================================================
# UI: BADGES (to match your screenshot)
# =========================================================
def badge(text: str, tone: str = "green"):
    colors = {
        "green": ("#1f7a1f", "#E7F6EA"),
        "yellow": ("#8a6a12", "#FBF3D0"),
        "red": ("#a11a1a", "#FDE2E2"),
        "blue": ("#1d4ed8", "#E8F0FF"),
        "gray": ("#374151", "#F3F4F6"),
    }
    fg, bg = colors.get(tone, colors["gray"])
    st.markdown(
        f"""
        <span style="
            display:inline-block;
            padding:10px 14px;
            border-radius:999px;
            background:{bg};
            border:1px solid rgba(0,0,0,0.08);
            color:{fg};
            font-weight:700;
            font-size:0.95rem;
            margin-right:10px;
            margin-bottom:8px;">
            {text}
        </span>
        """,
        unsafe_allow_html=True
    )


def tone_for_gate(gate: str) -> str:
    if gate == "PASS":
        return "green"
    if gate == "AT RISK":
        return "yellow"
    if gate == "FAIL":
        return "red"
    return "gray"


def diagnostics_panel(diag: Dict[str, Any]):
    with st.container(border=True):
        st.subheader("Diagnostics")
        st.caption("This helps you understand if the PDF text was actually readable.")

        q = (diag or {}).get("extraction_quality", "Unknown")
        tone = "green" if q == "Excellent" else ("yellow" if q in ["Good", "Fair"] else ("red" if q == "Poor" else "gray"))
        badge(f"Extraction quality: {q}", tone=tone)

        st.write(f"**File Type:** {diag.get('file_type', 'pdf')}")
        st.write(f"**Pages:** {diag.get('pages', 0)}    **Pages with Text:** {diag.get('pages_with_text', 0)}")
        st.write(f"**Characters Extracted:** {diag.get('characters_extracted', 0)}")
        st.write(f"**Likely Scanned:** {diag.get('likely_scanned', 'No')}")


def compliance_kpi_panel(kpis: Dict[str, Any]):
    with st.container(border=True):
        st.subheader("Compliance KPI")
        badge(f"Compliance: {int((kpis.get('completion', 0.0))*100)}%", "green" if kpis.get("completion", 0) >= 0.8 else ("yellow" if kpis.get("completion", 0) >= 0.4 else "red"))
        badge(f"Pass: {kpis.get('pass', 0)}", "green")
        badge(f"Fail: {kpis.get('fail', 0)}", "red" if kpis.get("fail", 0) > 0 else "green")
        badge(f"Unknown: {kpis.get('unknown', 0)}", "yellow" if kpis.get("unknown", 0) > 0 else "green")
        badge(f"Gate: {kpis.get('gate', 'WAITING')}", tone_for_gate(kpis.get("gate", "WAITING")))
        badge(f"Missing critical fields: {kpis.get('missing_critical', 0)}", "yellow" if kpis.get("missing_critical", 0) > 0 else "green")


# =========================================================
# AI (Modular, env-based; no extra requirements needed)
# =========================================================
def call_ai(prompt: str, system: str = "") -> Tuple[str, Optional[str]]:
    """
    Uses OpenAI Responses API if OPENAI_API_KEY is set.
    Returns (text, error_message_if_any).
    """
    api_key = st.secrets.get("OPENAI_API_KEY", None) if hasattr(st, "secrets") else None
    if not api_key:
        api_key = st.session_state.get("_OPENAI_API_KEY_OVERRIDE") or None
    if not api_key:
        return ("", "AI is not connected (missing OPENAI_API_KEY).")

    model = st.secrets.get("PATHAI_MODEL", "gpt-4.1-mini") if hasattr(st, "secrets") else "gpt-4.1-mini"

    payload = {
        "model": model,
        "input": [
            {"role": "system", "content": system or "You are a senior proposal writer and compliance expert."},
            {"role": "user", "content": prompt}
        ]
    }

    try:
        data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(
            "https://api.openai.com/v1/responses",
            data=data,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=90) as resp:
            raw = resp.read().decode("utf-8")
            obj = json.loads(raw)

        # Responses API: try the common paths
        text = ""
        if isinstance(obj, dict):
            # "output" is an array of content blocks
            out = obj.get("output", [])
            # Prefer any message content
            for item in out:
                if item.get("type") == "message":
                    content = item.get("content", [])
                    for c in content:
                        if c.get("type") in ["output_text", "text"]:
                            text += c.get("text", "")
            if not text:
                # fallback
                text = obj.get("output_text", "") or ""
        text = (text or "").strip()
        return (text, None if text else "AI returned empty output.")
    except Exception as e:
        return ("", f"AI request failed: {e}")


# =========================================================
# DOCX GENERATION (cover page + cover letter + proposal)
# =========================================================
def b64_from_uploaded_file(uploaded) -> str:
    if not uploaded:
        return ""
    try:
        content = uploaded.getvalue()
        return base64.b64encode(content).decode("utf-8")
    except Exception:
        return ""


def build_proposal_package_docx(
    company: Dict[str, Any],
    intake: Dict[str, Any],
    items: Dict[str, Item],
    cover_letter_text: str,
    proposal_text: str
) -> bytes:
    doc = docx.Document()

    # Cover page
    contract_title = clean_text(company.get("contract_title", "")) or "Solicitation"
    company_name = clean_text(company.get("legal_name", "")) or "Company Name"

    # Logo
    logo_b64 = company.get("logo_b64", "")
    if logo_b64:
        try:
            img_bytes = base64.b64decode(logo_b64)
            img_stream = io.BytesIO(img_bytes)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(2.0))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(company_name)
    r.bold = True
    r.font.size = docx.shared.Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(contract_title)
    r.bold = True
    r.font.size = docx.shared.Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated by {APP_NAME} • {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_page_break()

    # Cover letter
    h = doc.add_heading("Cover Letter", level=1)
    for line in (cover_letter_text or "").splitlines():
        doc.add_paragraph(line)

    # Signature block (required)
    doc.add_paragraph("")
    sig = doc.add_paragraph()
    sig.add_run("Sincerely,").bold = True
    doc.add_paragraph("")
    doc.add_paragraph(company.get("primary_contact", "") or "Authorized Representative")
    doc.add_paragraph(company_name)
    if company.get("email"):
        doc.add_paragraph(company["email"])
    if company.get("phone"):
        doc.add_paragraph(company["phone"])

    doc.add_page_break()

    # Proposal body
    doc.add_heading("Proposal", level=1)
    for line in (proposal_text or "").splitlines():
        doc.add_paragraph(line)

    # Appendix: Compliance snapshot (actionable only)
    actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
    if actionable:
        doc.add_page_break()
        doc.add_heading("Appendix: Compliance Snapshot", level=1)
        for it in sorted(actionable, key=lambda x: (-priority_score(x), -x.confidence)):
            status = it.status.upper()
            doc.add_paragraph(f"- [{status}] {it.normalized_text} (Source: {it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'})")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# =========================================================
# EXPORTS (XLSX + PDF + ZIP)
# =========================================================
def style_xlsx_header(ws, row=1):
    fill = PatternFill("solid", fgColor="EEF2FF")
    for cell in ws[row]:
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.alignment = Alignment(vertical="center", wrap_text=True)


def autosize_ws(ws, max_width=70):
    for col in ws.columns:
        max_len = 10
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, min(max_width, len(str(cell.value))))
        ws.column_dimensions[col_letter].width = max_len + 2


def export_xlsx_compliance_matrix(items: List[Item]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Compliance Matrix"

    ws.append([
        "id", "section", "mapped_section", "normalized_text",
        "gating_label", "confidence", "status", "done",
        "source", "bucket", "critical", "source_file", "page_number", "notes"
    ])

    for it in items:
        ws.append([
            it.id, it.section, it.mapped_section, it.normalized_text,
            it.gating_label, float(it.confidence), it.status, bool(it.done),
            it.source, it.bucket, bool(it.is_critical),
            it.source_file, it.page_number if it.page_number is not None else "",
            it.notes
        ])

    style_xlsx_header(ws)
    autosize_ws(ws)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def export_xlsx_task_list(items: List[Item]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Task List"

    actionable = [i for i in items if i.gating_label == GL_ACTIONABLE]
    actionable = sorted(actionable, key=lambda x: (-priority_score(x), -x.confidence))

    ws.append(["done", "priority_score", "bucket", "task", "mapped_section", "confidence", "status", "source_file", "page", "id", "critical"])
    for it in actionable:
        ws.append([
            "YES" if it.done else "NO",
            int(priority_score(it)),
            it.bucket,
            it.normalized_text,
            it.mapped_section,
            int(it.confidence * 100),
            it.status,
            it.source_file,
            it.page_number if it.page_number is not None else "",
            it.id,
            "YES" if it.is_critical else "NO"
        ])

    style_xlsx_header(ws)
    autosize_ws(ws)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def export_xlsx_gap_report(items: List[Item]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Gap Report"

    actionable = [i for i in items if i.gating_label == GL_ACTIONABLE]
    gaps = [i for i in actionable if not i.done and i.status != S_PASS]
    gaps = sorted(gaps, key=lambda x: (-priority_score(x), -x.confidence))

    ws.append(["gap", "priority_score", "bucket", "task", "mapped_section", "confidence", "status", "source_file", "page", "id", "critical"])
    for it in gaps:
        ws.append([
            "OPEN",
            int(priority_score(it)),
            it.bucket,
            it.normalized_text,
            it.mapped_section,
            int(it.confidence * 100),
            it.status,
            it.source_file,
            it.page_number if it.page_number is not None else "",
            it.id,
            "YES" if it.is_critical else "NO"
        ])

    style_xlsx_header(ws)
    autosize_ws(ws)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def pdf_write_wrapped(c: canvas.Canvas, x: int, y: int, text: str, max_chars: int = 100, line_height: int = 14) -> int:
    words = text.split()
    line = []
    cur = 0
    for w in words:
        if cur + len(w) + 1 > max_chars:
            c.drawString(x, y, " ".join(line))
            y -= line_height
            line = [w]
            cur = len(w)
        else:
            line.append(w)
            cur += len(w) + 1
    if line:
        c.drawString(x, y, " ".join(line))
        y -= line_height
    return y


def export_pdf_submission_checklist(items: List[Item], run_name: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    y = height - 54
    c.setFont("Helvetica-Bold", 16)
    c.drawString(48, y, f"{APP_NAME} — Submission Checklist")
    y -= 18
    c.setFont("Helvetica", 10)
    c.drawString(48, y, f"Run: {run_name}")
    y -= 14
    c.drawString(48, y, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    y -= 22

    actionable = [i for i in items if i.gating_label == GL_ACTIONABLE]
    grouped: Dict[str, List[Item]] = {}
    for it in actionable:
        grouped.setdefault(it.bucket, []).append(it)

    for bucket in ["Submission & Format", "Required Forms & Registrations", "Attachments/Exhibits", "Technical Requirements", "Pricing & Cost", "Past Performance", "Other"]:
        if bucket not in grouped:
            continue
        if y < 100:
            c.showPage()
            y = height - 54
        c.setFont("Helvetica-Bold", 12)
        c.drawString(48, y, bucket)
        y -= 18
        c.setFont("Helvetica", 10)
        for it in grouped[bucket]:
            if y < 80:
                c.showPage()
                y = height - 54
                c.setFont("Helvetica", 10)
            box = "[x]" if (it.done or it.status == S_PASS) else "[ ]"
            prefix = "CRITICAL: " if it.is_critical else ""
            src = f" ({it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'})"
            line = f"{box} {prefix}{it.normalized_text}{src}"
            y = pdf_write_wrapped(c, 58, y, line, max_chars=108)
            y -= 2

    c.showPage()
    c.save()
    return buf.getvalue()


def export_pdf_gate_report(items: List[Item], run_name: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    k = compute_kpis({i.id: i for i in items})
    y = height - 54
    c.setFont("Helvetica-Bold", 16)
    c.drawString(48, y, f"{APP_NAME} — Gate Report")
    y -= 18
    c.setFont("Helvetica", 10)
    c.drawString(48, y, f"Run: {run_name}")
    y -= 14
    c.drawString(48, y, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    y -= 20

    c.setFont("Helvetica-Bold", 12)
    c.drawString(48, y, "Summary")
    y -= 16
    c.setFont("Helvetica", 10)
    c.drawString(48, y, f"Gate: {k['gate']}")
    y -= 14
    c.drawString(48, y, f"Compliance: {int(k['completion']*100)}%")
    y -= 14
    c.drawString(48, y, f"Pass/Fail/Unknown: {k['pass']}/{k['fail']}/{k['unknown']}")
    y -= 14
    c.drawString(48, y, f"Missing critical fields: {k['missing_critical']}")
    y -= 24

    c.setFont("Helvetica-Bold", 12)
    c.drawString(48, y, "Top open gaps (actionable not yet done)")
    y -= 18
    c.setFont("Helvetica", 10)

    actionable = [i for i in items if i.gating_label == GL_ACTIONABLE]
    gaps = [i for i in actionable if (not i.done and i.status != S_PASS)]
    gaps = sorted(gaps, key=lambda x: (-priority_score(x), -x.confidence))

    if not gaps:
        c.drawString(48, y, "✅ No open gaps detected in actionable items.")
        y -= 14
    else:
        for it in gaps[:60]:
            if y < 80:
                c.showPage()
                y = height - 54
                c.setFont("Helvetica", 10)
            prefix = "CRITICAL: " if it.is_critical else ""
            src = f" — {it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}"
            y = pdf_write_wrapped(c, 52, y, f"• {prefix}{it.normalized_text}{src}", max_chars=112)

    c.showPage()
    c.save()
    return buf.getvalue()


def export_pdf_compliance_summary(items: List[Item], run_name: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    k = compute_kpis({i.id: i for i in items})
    state = "ON TRACK" if k["completion"] >= 0.80 and k["gate"] != "FAIL" else ("IN PROGRESS" if k["completion"] >= 0.40 else "BLOCKED")

    y = height - 54
    c.setFont("Helvetica-Bold", 16)
    c.drawString(48, y, f"{APP_NAME} — Compliance Summary")
    y -= 18
    c.setFont("Helvetica", 10)
    c.drawString(48, y, f"Run: {run_name}")
    y -= 14
    c.drawString(48, y, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    y -= 24

    c.setFont("Helvetica-Bold", 12)
    c.drawString(48, y, "Overall readiness")
    y -= 16
    c.setFont("Helvetica", 10)
    c.drawString(48, y, f"Status: {state}")
    y -= 14
    c.drawString(48, y, f"Gate: {k['gate']}")
    y -= 14
    c.drawString(48, y, f"Compliance: {int(k['completion']*100)}%")
    y -= 22

    c.setFont("Helvetica-Bold", 12)
    c.drawString(48, y, "Counts")
    y -= 16
    c.setFont("Helvetica", 10)
    c.drawString(48, y, f"Actionable total: {k['actionable_total']}")
    y -= 14
    c.drawString(48, y, f"Pass: {k['pass']}   Fail: {k['fail']}   Unknown: {k['unknown']}")
    y -= 14
    c.drawString(48, y, f"Missing critical: {k['missing_critical']}")
    y -= 22

    c.showPage()
    c.save()
    return buf.getvalue()


def export_zip_package(run_name: str, pdf1: bytes, pdf2: bytes, pdf3: bytes, xlsx1: bytes, xlsx2: bytes, xlsx3: bytes) -> bytes:
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("01_Submission_Checklist.pdf", pdf1)
        z.writestr("02_Gate_Report.pdf", pdf2)
        z.writestr("03_Compliance_Summary.pdf", pdf3)
        z.writestr("04_Compliance_Matrix.xlsx", xlsx1)
        z.writestr("05_Task_List.xlsx", xlsx2)
        z.writestr("06_Gap_Report.xlsx", xlsx3)
        z.writestr("README.txt", f"{APP_NAME} export package for run: {run_name}\nIncludes traceability (file + page) where available.\n")
    return zbuf.getvalue()


# =========================================================
# MAIN UI SECTIONS
# =========================================================
def header():
    st.title(APP_NAME)
    st.caption("You’re now on the Path to success.")
    st.caption(f"{APP_VERSION}")


def render_locked(message: str):
    st.info(f"🔒 {message}")


def reference_drawer(items: Dict[str, Item]):
    info_items = [i for i in items.values() if i.gating_label == GL_INFORMATIONAL]
    if not info_items:
        return
    with st.expander("Reference (Informational) — hidden by default", expanded=False):
        st.caption("These items are stored for context and traceability. They are not tasks.")
        q = st.text_input("Search reference", "")
        srcs = sorted(list(set(i.source_file for i in info_items)))
        src_filter = st.selectbox("Source file", ["All"] + srcs)
        filtered = info_items
        if q.strip():
            filtered = [i for i in filtered if q.lower() in i.normalized_text.lower()]
        if src_filter != "All":
            filtered = [i for i in filtered if i.source_file == src_filter]

        for it in filtered[:300]:
            src = f"{it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}"
            st.write(f"• **{it.section}** — {it.normalized_text}")
            st.caption(f"Source: {src}")


def render_task_list_home(run_id: str, run_meta: Dict[str, Any], items: Dict[str, Item]):
    intake = run_meta.get("intake", {}) or {}
    workflow = run_meta.get("workflow", {}) or {}
    company = run_meta.get("company", {}) or {}

    # --- Home: upload + paste + analyze (exact intent: first page)
    left, right = st.columns([0.62, 0.38], gap="large")

    with left:
        with st.container(border=True):
            st.subheader("Upload your RFP")
            pdfs = st.file_uploader("Upload PDF(s)", type=["pdf"], accept_multiple_files=True, key="home_pdfs")

            st.subheader("Or add RFP text")
            pasted = st.text_area("Paste text here (optional)", value=intake.get("sol_text", ""), height=160, key="home_text")

            include_reference = st.checkbox("Keep reference items (recommended)", value=True, key="home_keep_ref")

            analyze = st.button("Analyze", type="primary", use_container_width=True)

            if analyze:
                pages_data: List[Dict[str, Any]] = []
                sources: List[str] = []
                diag: Dict[str, Any] = {
                    "file_type": "pdf",
                    "pages": 0,
                    "pages_with_text": 0,
                    "characters_extracted": 0,
                    "likely_scanned": "No",
                    "extraction_quality": "Unknown"
                }

                if pdfs:
                    with st.spinner("Reading PDFs…"):
                        pages_data, sources, diag = extract_pages_from_pdfs(pdfs)

                # Combine text for critical detection
                combined_text_parts = []
                for p in pages_data:
                    combined_text_parts.append(p.get("text", ""))
                if pasted and pasted.strip():
                    combined_text_parts.append(pasted)

                combined_text = "\n".join(combined_text_parts)
                combined_text = combined_text or ""
                if not combined_text.strip():
                    st.warning("Upload at least one PDF or paste RFP text.")
                    st.stop()

                with st.spinner("Building your task list…"):
                    built = build_items_from_sources(combined_text, pages_data, pasted, include_reference)
                    built = dedupe_items(built)

                    delete_all_items(run_id)
                    for it in built:
                        upsert_item(run_id, it)

                    # Update run meta
                    intake["sol_text"] = pasted or ""
                    intake["sources"] = sources + (["Pasted text"] if pasted and pasted.strip() else [])
                    intake["diagnostics"] = diag
                    workflow["intake_complete"] = True

                    save_run_meta(run_id, intake=intake, company=company, workflow=workflow)

                st.success(f"Done. Built {len(built)} relevant items.")
                st.rerun()

        # --- The product: task list
        st.markdown("")
        st.subheader("Task List")
        st.caption("Only ACTIONABLE tasks show here.")

        actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
        if not actionable:
            st.info("Upload an RFP above and hit Analyze to generate your task list.")
            return

        # Top priorities + Blockers panels
        tp = top_priorities(items)
        bl = blockers(items)

        if bl:
            with st.container(border=True):
                st.subheader("Blockers")
                st.caption("These items can stop you from being compliant. Handle them first.")
                for it in bl:
                    cols = st.columns([0.08, 0.68, 0.24])
                    with cols[0]:
                        new_done = st.checkbox("", value=it.done, key=f"blk_done_{it.id}")
                    with cols[1]:
                        st.markdown(f"🚩 **{it.normalized_text}**")
                        src = f"{it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}"
                        with st.expander("View source", expanded=False):
                            st.write(f"**Source:** {src}")
                            st.write(it.source_snippet or it.raw_text)
                    with cols[2]:
                        new_status = st.selectbox("Status", [S_UNKNOWN, S_PASS, S_FAIL],
                                                  index=[S_UNKNOWN, S_PASS, S_FAIL].index(it.status),
                                                  key=f"blk_status_{it.id}")
                    it.done = bool(new_done)
                    it.status = new_status
                    upsert_item(run_id, it)

        if tp:
            with st.container(border=True):
                st.subheader("Top Priorities")
                st.caption("If you do nothing else, do these next.")
                for it in tp:
                    cols = st.columns([0.08, 0.68, 0.24])
                    with cols[0]:
                        new_done = st.checkbox("", value=it.done, key=f"tp_done_{it.id}")
                    with cols[1]:
                        st.markdown(f"**{it.normalized_text}**")
                        src = f"{it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}"
                        with st.expander("View source", expanded=False):
                            st.write(f"**Source:** {src}")
                            st.write(it.source_snippet or it.raw_text)
                    with cols[2]:
                        new_status = st.selectbox("Status", [S_UNKNOWN, S_PASS, S_FAIL],
                                                  index=[S_UNKNOWN, S_PASS, S_FAIL].index(it.status),
                                                  key=f"tp_status_{it.id}")
                    it.done = bool(new_done)
                    it.status = new_status
                    upsert_item(run_id, it)

        # Bucketed remainder
        st.markdown("")
        grouped: Dict[str, List[Item]] = {}
        for it in actionable:
            grouped.setdefault(it.bucket, []).append(it)

        for bucket in ["Submission & Format", "Required Forms & Registrations", "Attachments/Exhibits", "Technical Requirements", "Pricing & Cost", "Past Performance", "Other"]:
            if bucket not in grouped:
                continue
            bucket_items = sorted(grouped[bucket], key=lambda x: (-priority_score(x), -x.confidence))
            expanded = bucket in ["Submission & Format", "Required Forms & Registrations"]
            with st.expander(f"{bucket} ({len(bucket_items)})", expanded=expanded):
                for it in bucket_items:
                    cols = st.columns([0.08, 0.68, 0.24])
                    with cols[0]:
                        new_done = st.checkbox("", value=it.done, key=f"done_{it.id}")
                    with cols[1]:
                        prefix = "🚩 " if it.is_critical else ""
                        st.markdown(f"{prefix}**{it.normalized_text}**")
                        src = f"{it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}"
                        with st.expander("View source", expanded=False):
                            st.write(f"**Source:** {src}")
                            st.write(it.source_snippet or it.raw_text)
                    with cols[2]:
                        new_status = st.selectbox("Status", [S_UNKNOWN, S_PASS, S_FAIL],
                                                  index=[S_UNKNOWN, S_PASS, S_FAIL].index(it.status),
                                                  key=f"status_{it.id}")
                    it.done = bool(new_done)
                    it.status = new_status
                    upsert_item(run_id, it)

        reference_drawer(items)

    with right:
        # Diagnostics + KPI like screenshot
        diag = (intake or {}).get("diagnostics", {}) or {}
        kpis = compute_kpis(items)

        if diag:
            diagnostics_panel(diag)
        else:
            diagnostics_panel({
                "file_type": "pdf",
                "pages": 0,
                "pages_with_text": 0,
                "characters_extracted": 0,
                "likely_scanned": "No",
                "extraction_quality": "Unknown"
            })

        compliance_kpi_panel(kpis)

        st.markdown("")
        with st.container(border=True):
            st.subheader("Run details")
            st.write(f"**Run:** {run_meta.get('name', '')}")
            if company.get("contract_title"):
                st.write(f"**Contract:** {company.get('contract_title')}")
            if company.get("legal_name"):
                st.write(f"**Company:** {company.get('legal_name')}")


def render_intake(run_id: str, run_meta: Dict[str, Any], items: Dict[str, Item]):
    st.subheader("Intake")
    st.caption("Use Home (Task List) to upload and analyze. This tab is for review only.")

    intake = run_meta.get("intake", {}) or {}
    st.write("**Sources**")
    for s in (intake.get("sources", []) or []):
        st.write(f"• {s}")

    diag = intake.get("diagnostics", {}) or {}
    diagnostics_panel(diag if diag else {
        "file_type": "pdf",
        "pages": 0,
        "pages_with_text": 0,
        "characters_extracted": 0,
        "likely_scanned": "No",
        "extraction_quality": "Unknown"
    })


def render_company(run_id: str, run_meta: Dict[str, Any], items: Dict[str, Item]):
    st.subheader("Company")
    st.caption("Save your profile once. Path.ai uses this to auto-resolve and draft better.")

    company = run_meta.get("company", {}) or {}
    workflow = run_meta.get("workflow", {}) or {}

    with st.container(border=True):
        st.subheader("Cover page")
        company["contract_title"] = st.text_input("Contract / Solicitation title", value=company.get("contract_title", ""))
        logo = st.file_uploader("Upload company logo (PNG/JPG)", type=["png", "jpg", "jpeg"], key="logo_uploader")
        if logo is not None:
            company["logo_b64"] = b64_from_uploaded_file(logo)

        if company.get("logo_b64"):
            try:
                st.image(base64.b64decode(company["logo_b64"]), width=180)
            except Exception:
                pass

    st.markdown("")
    with st.container(border=True):
        st.subheader("Core company info")
        company["legal_name"] = st.text_input("Legal company name", value=company.get("legal_name", ""))
        company["uei"] = st.text_input("UEI", value=company.get("uei", ""))
        company["cage"] = st.text_input("CAGE", value=company.get("cage", ""))
        company["naics"] = st.text_input("Primary NAICS", value=company.get("naics", ""))
        company["address"] = st.text_area("Company address", value=company.get("address", ""), height=90)
        company["primary_contact"] = st.text_input("Primary contact", value=company.get("primary_contact", ""))
        company["email"] = st.text_input("Email", value=company.get("email", ""))
        company["phone"] = st.text_input("Phone", value=company.get("phone", ""))

    st.markdown("")
    with st.container(border=True):
        st.subheader("Registrations & certifications")
        company["sam_registered"] = st.checkbox("Registered in SAM.gov", value=bool(company.get("sam_registered", False)))
        cert_options = ["SDVOSB", "VetCert", "8(a)", "HUBZone", "WOSB"]
        company["certs"] = st.multiselect("Certifications (optional)", options=cert_options, default=company.get("certs", []))

    if st.button("Save Company", type="primary"):
        if company.get("legal_name", "").strip() and company.get("primary_contact", "").strip():
            workflow["company_complete"] = True

        save_run_meta(run_id, intake=run_meta.get("intake", {}), company=company, workflow=workflow)

        # Auto-resolve based on company
        items_latest = load_items(run_id)
        updated, changed = company_auto_resolve(items_latest, company)
        if changed > 0:
            for it in updated.values():
                upsert_item(run_id, it)
        st.success(f"Saved. Auto-resolved {changed} item(s) from your company profile.")


def render_compliance(run_id: str, run_meta: Dict[str, Any], items: Dict[str, Item]):
    st.subheader("Compliance")
    st.caption("Mark actionable requirements Pass/Fail/Unknown. This powers gate and exports.")

    actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
    if not actionable:
        st.info("No actionable items yet. Analyze an RFP on the Task List page.")
        return

    df = pd.DataFrame([{
        "Bucket": it.bucket,
        "Requirement": it.normalized_text,
        "Critical": "Yes" if it.is_critical else "",
        "Status": it.status,
        "Done": "Yes" if it.done else "",
        "Source": f"{it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}",
        "Confidence": int(it.confidence * 100)
    } for it in sorted(actionable, key=lambda x: (-priority_score(x), -x.confidence))])

    st.dataframe(df, use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("Gate Check")
    if st.button("Run Gate Check"):
        k = compute_kpis(items)
        if k["gate"] == "PASS":
            st.success("Gate: PASS — ready to proceed.")
        elif k["gate"] == "AT RISK":
            st.warning("Gate: AT RISK — unknown or missing critical fields remain.")
        else:
            st.error("Gate: FAIL — resolve failed items before proceeding.")

    # Auto-unlock compliance when far enough along
    workflow = run_meta.get("workflow", {}) or {}
    k = compute_kpis(items)
    if k["completion"] >= 0.60 and k["gate"] != "FAIL":
        if not workflow.get("compliance_complete", False):
            workflow["compliance_complete"] = True
            save_run_meta(run_id, intake=run_meta.get("intake", {}), company=run_meta.get("company", {}), workflow=workflow)
            st.success("Unlocked: Draft step.")


def render_draft(run_id: str, run_meta: Dict[str, Any], items: Dict[str, Item]):
    st.subheader("Proposal Generator")
    st.caption("Generate a compliant, competitive proposal draft using the RFP + your company profile + AI.")

    intake = run_meta.get("intake", {}) or {}
    company = run_meta.get("company", {}) or {}
    workflow = run_meta.get("workflow", {}) or {}

    actionable = [i for i in items.values() if i.gating_label == GL_ACTIONABLE]
    info_items = [i for i in items.values() if i.gating_label == GL_INFORMATIONAL]

    if not actionable and not info_items:
        st.info("Analyze an RFP first (Task List → Analyze).")
        return

    # Optional: allow key override inside app (for testing)
    with st.expander("AI connection (optional)", expanded=False):
        st.caption("If you don’t have secrets set yet, you can paste a key here for testing.")
        st.session_state["_OPENAI_API_KEY_OVERRIDE"] = st.text_input("OPENAI_API_KEY", type="password", value=st.session_state.get("_OPENAI_API_KEY_OVERRIDE", ""))

    contract_title = company.get("contract_title", "") or "Solicitation"
    company_name = company.get("legal_name", "") or "Company"

    # Build RFP context (compressed)
    top_reqs = sorted(actionable, key=lambda x: (-priority_score(x), -x.confidence))[:35]
    req_lines = []
    for it in top_reqs:
        src = f"{it.source_file}{'' if it.page_number is None else f' p.{it.page_number}'}"
        req_lines.append(f"- {it.normalized_text} (Status: {it.status}, Done: {it.done}, Source: {src})")
    req_block = "\n".join(req_lines)

    company_block = "\n".join([
        f"Company: {company.get('legal_name','')}",
        f"UEI: {company.get('uei','')}",
        f"CAGE: {company.get('cage','')}",
        f"NAICS: {company.get('naics','')}",
        f"Address: {company.get('address','')}",
        f"Primary Contact: {company.get('primary_contact','')}",
        f"Email: {company.get('email','')}",
        f"Phone: {company.get('phone','')}",
        f"SAM Registered: {company.get('sam_registered', False)}",
        f"Certifications: {', '.join(company.get('certs', []) or [])}"
    ])

    # UI controls
    colA, colB = st.columns([0.55, 0.45], gap="large")
    with colA:
        with st.container(border=True):
            st.subheader("Generate proposal draft")
            proposal_style = st.selectbox("Proposal style", ["Clear & compliant", "Executive & persuasive", "Technical & detailed"], index=0)
            length = st.select_slider("Length", options=["Short", "Medium", "Long"], value="Medium")
            focus = st.multiselect("Focus", options=["Compliance-first", "Differentiators", "Risk mitigation", "Past performance framing", "Pricing narrative (non-pricing)"], default=["Compliance-first", "Differentiators"])

            gen_proposal = st.button("Generate Proposal Draft", type="primary", use_container_width=True)

            if gen_proposal:
                system = (
                    "You are a senior federal proposal writer, compliance manager, and capture strategist. "
                    "Your job is to produce a proposal that is maximally compliant and highly competitive. "
                    "You must be specific, structured, and aligned to what the solicitation requires. "
                    "If details are missing, write clear placeholders and list exactly what must be confirmed."
                )
                prompt = f"""
Create a proposal draft for this solicitation.

Solicitation title: {contract_title}

Company profile:
{company_block}

High-priority compliance requirements (actionable):
{req_block}

User preference:
- Style: {proposal_style}
- Length: {length}
- Focus: {", ".join(focus) if focus else "Compliance-first"}

Output requirements:
1) Use clear headings and sections.
2) Include a Compliance Mapping section that references how the draft addresses the actionable requirements.
3) Include a “Missing Inputs” list if anything is unknown or ambiguous.
4) Write in a confident, competitive tone without fluff.
5) Do not invent contract details like dates or exact submission portals. Use placeholders when needed.

Return only the proposal draft text.
"""
                with st.spinner("Generating proposal draft…"):
                    text, err = call_ai(prompt=prompt, system=system)
                drafts = intake.get("drafts", {}) or {}
                if err:
                    st.error(err)
                    st.info("Tip: Set OPENAI_API_KEY in Render environment variables (or paste one in the expander above).")
                    # fallback
                    text = f"""PROPOSAL DRAFT (TEMPLATE FALLBACK)

Solicitation: {contract_title}
Offeror: {company_name}

1. Executive Summary
- [Describe understanding of requirement]
- [Key differentiators]

2. Technical Approach
- [Approach aligned to PWS/SOW]
- [Staffing and management]
- [Quality control]

3. Compliance Mapping (Actionable)
{req_block if req_block else "- [No actionable requirements captured yet]"}

4. Past Performance
- [Insert relevant past performance references]

5. Risk Mitigation
- [Risks + mitigations]

6. Missing Inputs
- [List submission deadline, portal, format, page limits, required forms, etc. if missing]
"""
                drafts["proposal"] = text
                intake["drafts"] = drafts
                save_run_meta(run_id, intake=intake, company=company, workflow=workflow)
                st.success("Proposal draft generated.")
                st.rerun()

            drafts = intake.get("drafts", {}) or {}
            if drafts.get("proposal"):
                st.text_area("Proposal draft", value=drafts["proposal"], height=420)

    with colB:
        with st.container(border=True):
            st.subheader("Generate cover letter")
            tone = st.selectbox("Cover letter tone", ["Professional & direct", "Warm & confident", "Bold & competitive"], index=0)
            gen_cover = st.button("Generate Cover Letter", use_container_width=True)

            if gen_cover:
                system = (
                    "You are a senior federal proposal writer. "
                    "Write a cover letter that is compliant, confident, and tailored to the solicitation. "
                    "Use the company info provided. Include placeholders for unknown details. "
                    "End with a signature block."
                )
                prompt = f"""
Write a cover letter for this proposal submission.

Solicitation title: {contract_title}

Company profile:
{company_block}

Key submission/compliance items (high priority):
{req_block}

Tone: {tone}

Cover letter requirements:
- Mention the solicitation title.
- State the offeror’s intent to submit a compliant proposal.
- Briefly highlight differentiators (do not invent).
- Include a compliance-forward sentence (e.g., "We have reviewed the submission instructions and confirm our proposal is prepared to comply.")
- End with a signature block:
  Sincerely,
  <Name>
  <Title or Authorized Representative>
  <Company>
  <Email>
  <Phone>

Return only the cover letter text.
"""
                with st.spinner("Generating cover letter…"):
                    text, err = call_ai(prompt=prompt, system=system)

                drafts = intake.get("drafts", {}) or {}
                if err:
                    st.error(err)
                    text = f"""COVER LETTER (TEMPLATE FALLBACK)

Date: {datetime.utcnow().strftime('%Y-%m-%d')}

Subject: Proposal Submission — {contract_title}

Dear Contracting Officer,

{company_name} is pleased to submit our proposal in response to the solicitation titled "{contract_title}." We have reviewed the submission requirements and have prepared our response to comply with the instructions and required forms.

We bring disciplined execution, clear communication, and a compliance-first delivery approach. Where the solicitation requires confirmation of specific submission details (deadline, portal, format), we have flagged those items for verification and will finalize them prior to submission.

Thank you for your consideration.

Sincerely,

{company.get("primary_contact","Authorized Representative")}
{company_name}
{company.get("email","")}
{company.get("phone","")}
"""
                drafts["cover_letter"] = text
                intake["drafts"] = drafts
                save_run_meta(run_id, intake=intake, company=company, workflow=workflow)
                st.success("Cover letter generated.")
                st.rerun()

            drafts = intake.get("drafts", {}) or {}
            if drafts.get("cover_letter"):
                st.text_area("Cover letter", value=drafts["cover_letter"], height=320)

        st.markdown("")
        with st.container(border=True):
            st.subheader("Download package (DOCX)")
            drafts = intake.get("drafts", {}) or {}
            if not drafts.get("proposal") or not drafts.get("cover_letter"):
                st.info("Generate both the proposal draft and cover letter first.")
            else:
                docx_bytes = build_proposal_package_docx(
                    company=company,
                    intake=intake,
                    items=items,
                    cover_letter_text=drafts.get("cover_letter", ""),
                    proposal_text=drafts.get("proposal", "")
                )
                st.download_button(
                    "Download Proposal Package (DOCX)",
                    data=docx_bytes,
                    file_name="PathAI_Proposal_Package.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

    st.markdown("---")
    st.caption("When you’re happy with the draft, mark it ready to unlock Export.")
    ready = st.checkbox("Draft is ready", value=bool(workflow.get("draft_complete", False)))
    if ready != bool(workflow.get("draft_complete", False)):
        workflow["draft_complete"] = bool(ready)
        save_run_meta(run_id, intake=intake, company=company, workflow=workflow)
        st.success("Updated.")


def render_export(run_id: str, run_meta: Dict[str, Any], items: Dict[str, Item]):
    st.subheader("Export")
    st.caption("Exports include traceability (file + page) where available.")

    all_items = list(items.values())
    if not all_items:
        st.info("Nothing to export yet. Analyze an RFP first.")
        return

    run_name = run_meta.get("name", "My Proposal Run")

    xlsx_matrix = export_xlsx_compliance_matrix(all_items)
    xlsx_tasks = export_xlsx_task_list(all_items)
    xlsx_gaps = export_xlsx_gap_report(all_items)

    pdf_checklist = export_pdf_submission_checklist(all_items, run_name)
    pdf_gate = export_pdf_gate_report(all_items, run_name)
    pdf_summary = export_pdf_compliance_summary(all_items, run_name)

    zip_bytes = export_zip_package(run_name, pdf_checklist, pdf_gate, pdf_summary, xlsx_matrix, xlsx_tasks, xlsx_gaps)

    st.download_button("Download Submission Checklist (PDF)", data=pdf_checklist, file_name="Submission_Checklist.pdf", mime="application/pdf")
    st.download_button("Download Gate Report (PDF)", data=pdf_gate, file_name="Gate_Report.pdf", mime="application/pdf")
    st.download_button("Download Compliance Summary (PDF)", data=pdf_summary, file_name="Compliance_Summary.pdf", mime="application/pdf")

    st.download_button("Download Compliance Matrix (XLSX)", data=xlsx_matrix, file_name="Compliance_Matrix.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    st.download_button("Download Task List (XLSX)", data=xlsx_tasks, file_name="Task_List.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    st.download_button("Download Gap Report (XLSX)", data=xlsx_gaps, file_name="Gap_Report.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.download_button("Download FULL Export Package (ZIP)", data=zip_bytes, file_name="PathAI_Export_Package.zip", mime="application/zip")


# =========================================================
# STREAMLIT APP
# =========================================================
st.set_page_config(page_title=APP_NAME, layout="wide")

# Sidebar (you said you're ok with it)
with st.sidebar:
    st.markdown("## Proposal Runs")
    runs = list_runs()

    if "active_run_id" not in st.session_state:
        st.session_state.active_run_id = None

    if st.button("➕ New run", use_container_width=True):
        new_name = f"My Proposal Run ({datetime.utcnow().strftime('%Y-%m-%d')})"
        rid = create_run(new_name)
        st.session_state.active_run_id = rid
        st.success("Created.")

    run_options = ["(Select a run)"] + [f"{name} — {rid}" for rid, name, _ in runs]
    sel = st.selectbox("Select run", run_options, index=0)

    if sel != "(Select a run)":
        rid = sel.split("—")[-1].strip()
        st.session_state.active_run_id = rid

    st.markdown("---")
    st.caption("Home = Upload + Analyze + Task List.\n\nOnly ACTIONABLE items become tasks.")

run_id = st.session_state.active_run_id
if not run_id:
    st.title(APP_NAME)
    st.info("Create or select a proposal run from the sidebar to begin.")
    st.stop()

run_meta = load_run_meta(run_id)
items = load_items(run_id)

# header
header()

# auto-save workflow unlock updates
unlocks = workflow_unlocks(run_meta, items)
save_run_meta(run_id, intake=run_meta.get("intake", {}), company=run_meta.get("company", {}), workflow=run_meta.get("workflow", {}), run_name=run_meta.get("name", None))

tabs = st.tabs(PAGES)

with tabs[0]:
    # Refresh run meta/items each render for accurate KPIs
    run_meta = load_run_meta(run_id)
    items = load_items(run_id)
    render_task_list_home(run_id, run_meta, items)

with tabs[1]:
    run_meta = load_run_meta(run_id)
    items = load_items(run_id)
    render_intake(run_id, run_meta, items)

with tabs[2]:
    run_meta = load_run_meta(run_id)
    items = load_items(run_id)
    if not unlocks["Company"]:
        render_locked("Complete Analyze on the Task List page first.")
    else:
        render_company(run_id, run_meta, items)

with tabs[3]:
    run_meta = load_run_meta(run_id)
    items = load_items(run_id)
    unlocks = workflow_unlocks(run_meta, items)
    if not unlocks["Compliance"]:
        render_locked("Save Company info first.")
    else:
        render_compliance(run_id, run_meta, items)

with tabs[4]:
    run_meta = load_run_meta(run_id)
    items = load_items(run_id)
    unlocks = workflow_unlocks(run_meta, items)
    if not unlocks["Draft"]:
        render_locked("Complete enough Compliance first (about 60% and no FAIL).")
    else:
        render_draft(run_id, run_meta, items)

with tabs[5]:
    run_meta = load_run_meta(run_id)
    items = load_items(run_id)
    unlocks = workflow_unlocks(run_meta, items)
    if not unlocks["Export"]:
        render_locked("Mark Draft as ready to unlock Export.")
    else:
        render_export(run_id, run_meta, items)