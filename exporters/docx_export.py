from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Inches


def build_docx_bytes(
    *,
    rfp: Dict[str, Any],
    company: Dict[str, Any],
    cover_letter: str,
    proposal_body: str,
    logo_bytes: Optional[bytes] = None,
) -> bytes:
    doc = Document()

    if logo_bytes:
        try:
            doc.add_picture(BytesIO(logo_bytes), width=Inches(1.5))
        except Exception:
            pass

    doc.add_heading(company.get("name") or "Proposal Package", level=1)

    fn = (rfp.get("filename") or "").strip()
    if fn:
        doc.add_paragraph(f"RFP: {fn}")
    if rfp.get("due_date"):
        doc.add_paragraph(f"Due Date: {rfp.get('due_date')}")
    if rfp.get("submission_email"):
        doc.add_paragraph(f"Submission: {rfp.get('submission_email')}")

    doc.add_page_break()

    doc.add_heading("Cover Letter", level=1)
    for line in (cover_letter or "").split("\n"):
        doc.add_paragraph(line)

    doc.add_page_break()

    doc.add_heading("Proposal", level=1)
    for line in (proposal_body or "").split("\n"):
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()